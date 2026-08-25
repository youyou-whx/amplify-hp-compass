# -*- coding: utf-8 -*-
"""在 HP COMPASS.docx 的稳定性章节补充"可选模式"说明。"""
import sys

from docx.shared import Inches, Pt, RGBColor

SKILL_DIR = r"C:\Users\whx\.claude\skills\igem-conservation-reference"
sys.path.insert(0, SKILL_DIR)
from docx_template_helper import set_run_font

doc_path = r"C:\Users\whx\igem1\HP COMPASS.docx"
doc = __import__("docx").Document(doc_path)

# 找到"（五）LLM 层稳定性（实测数据）"标题段
target = None
for i, para in enumerate(doc.paragraphs):
    if "LLM 层稳定性" in para.text:
        target = i
        break

if target is None:
    print("未找到稳定性章节标题")
    sys.exit(1)

# 找到标题后的第一段正文（"大模型层调用 1/2 各运行两次..."），在其后插入说明段
insert_after = target
for j in range(target + 1, len(doc.paragraphs)):
    if doc.paragraphs[j].text.strip():
        insert_after = j
        break

new_para = doc.paragraphs[insert_after].insert_paragraph_before()
new_para.paragraph_format.first_line_indent = Inches(0.3)
new_para.paragraph_format.line_spacing_rule = 3  # WD_LINE_SPACING.ONE_POINT_FIVE
run = new_para.add_run(
    "稳定性检查为可选测试模式：日常解析默认单次运行，不重复调用；"
    "进行稳定性测试时（测试脚本加 --stability 参数，或网页勾选 Stability check），"
    "每条记录双运行并比较 15 个四梯度字段的一致率，采用第一次运行结果，两次原始 JSON 均存档。"
)
set_run_font(run, font_name="宋体", size=11, color=RGBColor(0x3D, 0x2B, 0x1F))

doc.save(doc_path)
print("已插入稳定性说明段（位于第 %d 段之后）" % insert_after)
