"""Export all white-box model parameters to a structured Word document.

Reads config.py and related modules, writes a formatted .docx.
Does NOT modify any existing files.

v3: Maturity FCE fully implemented — signal weights w_{i,j} and membership
    function parameters μ_{i,k}^{(j)} are now explicit, calibrated parameters.
"""

from __future__ import annotations

import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from hp_compass.config import (
    CATEGORIES, MODULE_THRESHOLDS, MODULE_CRITICALITY,
    J1_MATRIX, J2_MATRIX, J2ND_MATRIX,
    A1_WEIGHTS, A2_WEIGHTS, A2ND_WEIGHTS,
    QUANT_VECTOR, MEMBERSHIP_PARAMS,
    MATURITY_SIGNAL_WEIGHTS, MATURITY_MEMBERSHIP_PARAMS,
    EVIDENCE_KEYWORDS, EVIDENCE_BLOCKED,
    ACTION_KEYWORDS, RETURN_KEYWORDS,
    STAKEHOLDER_VALUE_KEYWORDS, STAKEHOLDER_TYPE_VALUES,
)
from hp_compass.maturity import (
    DIMENSION_ORDER, DIMENSION_LABELS, DIMENSION_SHORT, LEVEL_ANCHORS,
    DESIGN_MODULES, DESIGN_REFLECTION_KEYWORDS, DESIGN_ITERATION_KEYWORDS,
    REAL_WORLD_KEYWORDS, REAL_WORLD_STAKEHOLDER_TYPES, REAL_WORLD_MODULES,
    UNIQUE_STAKEHOLDER_TYPES, RISK_KEYWORDS, RISK_EVIDENCE_ARTIFACTS,
    RESPONSE_ACTION_KEYWORDS, LIMITATION_KEYWORDS,
    BOUNDARY_LANGUAGE, BOUNDARY_ARTIFACTS,
    MITIGATION_KEYWORDS,
    _DESIGN_KW_SATURATION, _SCENE_KW_SATURATION, _RISK_KW_SATURATION,
    _MITIGATION_KW_SATURATION, _BOUNDARY_KW_SATURATION,
    _LIMITATION_KW_DENOM, _ACTION_LEN_DENOM, _MODULE_COUNT_DENOM,
)
from hp_compass.graph_builder import (
    NODE_LEVELS, LEVEL_LABELS, THEME_PATTERNS, STRONG_TERMS, ENTITY_TERMS,
)
from hp_compass.status import STATUS_NAMES
from hp_compass.sensitivity import FACTOR_LABELS


# ═══════════════════════════════════════════════════════
#  Helper functions
# ═══════════════════════════════════════════════════════

def set_cell_shading(cell, color_hex: str):
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color_hex,
    })
    shading.append(shading_elm)


def add_table(doc, headers: list[str], rows: list[list[str]],
              col_widths: list[float] | None = None,
              header_color: str = "1B5E8A") -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            if i % 2 == 1:
                set_cell_shading(cell, "F5F0E8")
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)


def add_heading(doc, text: str, level: int = 1):
    return doc.add_heading(text, level=level)


def add_para(doc, text: str):
    return doc.add_paragraph(text)


# ═══════════════════════════════════════════════════════
#  Build document
# ═══════════════════════════════════════════════════════

def main():
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # ── Title ──
    title = doc.add_heading("HP Compass 模型白盒参数总表", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "AMPlify 2026 · iGEM Conservation 团队 · Human Practices 决策支持模型")
    add_para(doc, "本文档汇总模型中全部可配置的白盒参数——词集、权重、阈值、隶属函数参数、"
              "AHP判断矩阵、成熟度信号权重与隶属函数参数等。所有参数均可在 config.py 及对应模块源文件中直接查阅。")
    add_para(doc, "版本：v3.1（2026-07-15），对应模型规范 v3.1 — "
              "成熟度参数优化：隶属函数阈值下调~30%，零信号均匀先验，归一化参数放宽。")
    doc.add_paragraph("")

    # ════════════════════════════════════════
    #  1. Module Categories & Keywords
    # ════════════════════════════════════════
    add_heading(doc, "1. 模块分类与关键词集（Φ₂ 模糊隶属度分类）", 1)
    add_para(doc, "模型将 HP 反馈文本映射到 9 个项目模块。每模块有一个特征词集 K_c，"
              "命中密度 h_c = |M_c(T*)| / |K_c| 经升半梯形隶属函数转换为模糊隶属度 μ_c ∈ [0,1]。")

    for cat, terms in CATEGORIES.items():
        idx = list(CATEGORIES.keys()).index(cat) + 1
        add_heading(doc, f"1.{idx} {cat}（|K_c| = {len(terms)}）", 2)
        add_table(doc,
                  ["序号", "关键词", "序号", "关键词", "序号", "关键词"],
                  _chunk_list(terms, 3),
                  col_widths=[1.0, 5.0, 1.0, 5.0, 1.0, 5.0])

    # ════════════════════════════════════════
    #  2. Module Thresholds & Criticality
    # ════════════════════════════════════════
    add_heading(doc, "2. 模块隶属度阈值与关键性系数", 1)
    add_para(doc, "升半梯形隶属函数 μ_c(h_c; α_c, β_c)：h_c < α_c → 0, "
              "α_c ≤ h_c < β_c → (h_c−α_c)/(β_c−α_c), h_c ≥ β_c → 1。\n"
              "参数经 Delphi 法标定（3位iGEM领域专家，2轮问卷，CV < 0.15）。\n"
              "κ(c) 为模块关键性系数，用于优先级因子 F₃。")

    rows = []
    for cat in CATEGORIES:
        alpha, beta = MODULE_THRESHOLDS[cat]
        kappa = MODULE_CRITICALITY[cat]
        rows.append([cat, str(len(CATEGORIES[cat])), f"{alpha:.2f}", f"{beta:.2f}", f"{kappa:.2f}"])
    add_table(doc, ["模块", "|K_c|", "α_c（激活）", "β_c（饱和）", "κ(c) 关键性"],
              rows, col_widths=[3.5, 1.5, 2.8, 2.8, 2.8])

    # ════════════════════════════════════════
    #  3. AHP Judgment Matrices
    # ════════════════════════════════════════
    add_heading(doc, "3. AHP 判断矩阵与权重（Φ₄ 优先级）", 1)
    add_para(doc, "采用 Saaty 1-9 标度。3位专家独立填写，几何平均，两轮 Delphi 迭代至 CV < 0.15。"
              "方根法求解权重，CR < 0.10 通过一致性检验。RI₃=0.58。")

    add_heading(doc, "3.1 J₁：内部紧迫性 U₁ = (F₁, F₂, F₃)", 2)
    add_para(doc, "F₁（闭环缺口）比 F₂ 稍微重要（标度2），比 F₃ 稍微重要（标度3）；F₂ 与 F₃ 同等重要。")
    _add_matrix_table(doc, J1_MATRIX, ["F₁ 闭环缺口", "F₂ 跨模块影响", "F₃ 关键性"])
    add_para(doc, f"→ A₁ = ({A1_WEIGHTS[0]:.3f}, {A1_WEIGHTS[1]:.3f}, {A1_WEIGHTS[2]:.3f})  "
              f"λ_max=3.018, CI=0.009, CR=0.016 ✓")

    add_heading(doc, "3.2 J₂：外部约束 U₂ = (F₄, F₅, F₆)", 2)
    add_para(doc, "F₄（时间紧迫度）比 F₅ 和 F₆ 稍微重要（标度2）；F₅ 与 F₆ 同等重要。")
    _add_matrix_table(doc, J2_MATRIX, ["F₄ 时间紧迫度", "F₅ 证据不足度", "F₆ 利益相关者价值"])
    add_para(doc, f"→ A₂ = ({A2_WEIGHTS[0]:.3f}, {A2_WEIGHTS[1]:.3f}, {A2_WEIGHTS[2]:.3f})  "
              f"λ_max=3.001, CI=0.001, CR=0.001 ✓")

    add_heading(doc, "3.3 J：第二级综合（U₁ vs U₂）", 2)
    add_para(doc, "U₁（内部紧迫性）比 U₂（外部约束）稍微重要（标度2）。")
    _add_matrix_table(doc, J2ND_MATRIX, ["U₁ 内部紧迫性", "U₂ 外部约束"])
    add_para(doc, f"→ A = ({A2ND_WEIGHTS[0]:.3f}, {A2ND_WEIGHTS[1]:.3f})，二阶矩阵恒一致 ✓")

    # ════════════════════════════════════════
    #  4. FCE Membership & Quantification
    # ════════════════════════════════════════
    add_heading(doc, "4. FCE 评语集量化向量与隶属函数参数（Φ₄）", 1)

    add_heading(doc, "4.1 评语集与量化向量 C", 2)
    add_para(doc, "评语集 V = {v₁, v₂, v₃, v₄} = {低, 中, 高, 紧急}。重心法去模糊化：P = Σ bⱼ·cⱼ。")
    add_table(doc,
              ["参数", "v₁ 低", "v₂ 中", "v₃ 高", "v₄ 紧急"],
              [["cⱼ", "0.20", "0.45", "0.72", "0.95"],
               ["Δcⱼ（间距）", "—", "+0.25", "+0.27", "+0.23"]],
              col_widths=[3.0, 2.5, 2.5, 2.5, 2.5])

    add_heading(doc, "4.2 隶属函数形状参数 (a, b, c, d)", 2)
    add_para(doc, "四类梯形/半梯形隶属函数，参数经 Delphi 法标定（3位专家，2轮，CV < 0.15）。"
              "对每个因素 F_k ∈ [0,1]，计算其对四个评语等级的隶属度，构成 3×4 模糊评价矩阵 R。"
              "满足：(1) 覆盖性 Σⱼ r_{ij}(F) = 1；(2) 凸性；(3) 平滑过渡。")
    shape_descriptions = {
        "v1_low":    "降半梯形（低）r₁",
        "v2_mid":    "三角形（中）r₂",
        "v3_high":   "梯形（高）r₃",
        "v4_urgent": "升半梯形（紧急）r₄",
    }
    rows = []
    for key in ["v1_low", "v2_mid", "v3_high", "v4_urgent"]:
        a, b, c, d = MEMBERSHIP_PARAMS[key]
        rows.append([shape_descriptions[key], f"{a:.2f}", f"{b:.2f}", f"{c:.2f}", f"{d:.2f}"])
    add_table(doc, ["函数形状", "a", "b", "c", "d"], rows,
              col_widths=[5.0, 2.0, 2.0, 2.0, 2.0])

    # ════════════════════════════════════════
    #  5. Evidence & Action Keywords
    # ════════════════════════════════════════
    add_heading(doc, "5. 证据与行动关键词库（Φ₃ 闭环状态与证据强度）", 1)

    add_heading(doc, "5.1 证据关键词权重表", 2)
    add_para(doc, "用于计算 evidence_strength：对每个证据项 e_i，取匹配到的最高权重关键词分数；"
              "所有证据项取平均。未匹配任何关键词的证据项默认基准分 0.35。")
    rows = sorted(
        [[kw, f"{w:.2f}"] for kw, w in EVIDENCE_KEYWORDS.items()],
        key=lambda x: -float(x[1])
    )
    add_table(doc, ["证据关键词", "权重 σ(e)"], rows, col_widths=[8.0, 3.0])

    add_heading(doc, "5.2 证据屏蔽词", 2)
    add_para(doc, "包含以下计划性/未来式表述的文本不被视为有效证据：")
    add_para(doc, "、".join(sorted(EVIDENCE_BLOCKED)))

    add_heading(doc, "5.3 行动关键词", 2)
    add_para(doc, "用于提取 project_action 字段及维度5信号计算：")
    add_para(doc, "、".join(ACTION_KEYWORDS))

    add_heading(doc, "5.4 返回/闭环确认关键词", 2)
    add_para(doc, "用于检测 returned 状态（L4）：")
    add_para(doc, "【显式确认模式】" + "、".join([
        "已返回", "已经返回", "已回访", "已经回访", "获得二轮反馈",
        "完成二轮反馈", "完成了二轮反馈", "二轮反馈确认",
        "returned and confirmed", "second feedback confirmed",
    ]))
    add_para(doc, "【计划性模式（命中则跳过）】" + "、".join([
        "下一步应", "下一步", "计划", "仍需", "需要", "将", "应把", "后续",
        "Second Feedback", "反馈计划",
    ]))

    # ════════════════════════════════════════
    #  6. Stakeholder Value
    # ════════════════════════════════════════
    add_heading(doc, "6. 利益相关者价值参数（F₆ 计算）", 1)
    add_para(doc, "F₆ = max(0.50, max_{k⊑s} v(k))。底线 0.50 保证任何利益相关者的反馈都具有基础价值。")

    add_heading(doc, "6.1 文本关键词匹配价值", 2)
    rows = [[kw, f"{v:.2f}"] for kw, v in sorted(
        STAKEHOLDER_VALUE_KEYWORDS.items(), key=lambda x: -x[1]
    )]
    add_table(doc, ["角色关键词", "价值 v(k)"], rows, col_widths=[8.0, 3.0])

    add_heading(doc, "6.2 Stakeholder 类型基准价值", 2)
    rows = [[t, f"{v:.2f}"] for t, v in sorted(
        STAKEHOLDER_TYPE_VALUES.items(), key=lambda x: -x[1]
    )]
    add_table(doc, ["Stakeholder 类型", "基准价值"], rows, col_widths=[8.0, 3.0])

    # ════════════════════════════════════════
    #  7. Maturity Dimension Anchors
    # ════════════════════════════════════════
    add_heading(doc, "7. 成熟度六维锚定标准（Φ₅）", 1)
    add_para(doc, "每个维度 0-5 级。FCE 合成后采用最大隶属度原则判定等级"
              "（若 ∃k: μ_{i,k} > 0.5），同时计算级别特征值 m_i* = Σ(k·μ_k^γ)/Σ(μ_k^γ)，γ=2。")

    for dim_key in DIMENSION_ORDER:
        idx = DIMENSION_ORDER.index(dim_key) + 1
        full = DIMENSION_LABELS[dim_key]
        short = DIMENSION_SHORT[dim_key]
        add_heading(doc, f"7.{idx} {short}（{full}）", 2)
        anchors = LEVEL_ANCHORS[dim_key]
        rows = [[str(k), anchors[k]] for k in range(6)]
        add_table(doc, ["等级", "锚定描述"], rows, col_widths=[1.5, 12.0])

    # ════════════════════════════════════════
    #  8. Maturity Signal Keywords
    # ════════════════════════════════════════
    add_heading(doc, "8. 成熟度信号关键词库（维度1-6）", 1)
    add_para(doc, "以下关键词集用于计算各维度的信号变量 x_{i,j}（keyword density / count）。")

    sections = [
        ("8.1 设计相关模块（维度1）", sorted(DESIGN_MODULES)),
        ("8.2 设计反思关键词（维度1，饱和点=6）", DESIGN_REFLECTION_KEYWORDS),
        ("8.3 迭代关键词（维度1 高级信号，已纳入反思关键词集）", DESIGN_ITERATION_KEYWORDS),
        ("8.4 真实场景关键词（维度2，饱和点=6）", REAL_WORLD_KEYWORDS),
        ("8.5 真实场景 Stakeholder 类型（维度2）", sorted(REAL_WORLD_STAKEHOLDER_TYPES)),
        ("8.6 真实场景关联模块（维度2）", sorted(REAL_WORLD_MODULES)),
        ("8.7 独特/关键 Stakeholder 类型（维度3）", sorted(UNIQUE_STAKEHOLDER_TYPES)),
        ("8.8 风险关键词（维度4，饱和点=6）", RISK_KEYWORDS),
        ("8.9 风险证据产物（维度4）", RISK_EVIDENCE_ARTIFACTS),
        ("8.10 缓解措施关键词（维度4，饱和点=4）", MITIGATION_KEYWORDS),
        ("8.11 行动响应关键词（维度5，与 §5.3 互补）", RESPONSE_ACTION_KEYWORDS),
        ("8.12 局限性关键词（维度6，分母=8）", LIMITATION_KEYWORDS),
        ("8.13 边界语言（维度6 高级信号，饱和点=3）", BOUNDARY_LANGUAGE),
        ("8.14 边界文档产物（维度6）", BOUNDARY_ARTIFACTS),
    ]
    for title, items in sections:
        add_heading(doc, title, 2)
        add_para(doc, "、".join(items))

    # ════════════════════════════════════════
    #  9. Loop Status
    # ════════════════════════════════════════
    add_heading(doc, "9. 闭环状态机 L0-L4（Φ₃）", 1)
    add_para(doc, "判定采用最大达成原则：ℓ = max(𝟏_{|f|>0}, 2·𝟏_{|a|>0}, 3·𝟏_{|E|>0∧σₑ>0}, 4·𝟏_{r=1})。"
              "回访对话要么发生、要么未发生，属质变而非度变，保留离散判定。")
    rows = [
        ["L0", "L0_Recorded", "无反馈文本"],
        ["L1", "L1_Interpreted", "有反馈文本 或 有模块分类"],
        ["L2", "L2_Actioned", "project_action 非空"],
        ["L3", "L3_Evidenced", "evidence 非空 且 evidence_strength > 0"],
        ["L4", "L4_Returned", "returned = True"],
    ]
    add_table(doc, ["等级", "状态标识", "判定条件"], rows, col_widths=[1.8, 3.5, 8.0])

    # ════════════════════════════════════════
    #  10. Extractor Stakeholder Mapping
    # ════════════════════════════════════════
    add_heading(doc, "10. 文本提取器 Stakeholder 映射表（Φ₁）", 1)

    add_heading(doc, "10.1 已知 Stakeholder 精确映射", 2)
    add_para(doc, "从文件名或文本首部（前1800字符）匹配：")
    rows = [
        ["赵天意", "赵天意老师"],
        ["钱勋", "西北农林科技大学资环学院钱勋教授"],
        ["罗自卫", "西北农林科技大学罗自卫老师"],
        ["聂桓 / 哈工大", "哈尔滨工业大学生命科学与医学学部聂桓老师"],
        ["东北地区交流会", "iGEM 东北地区交流会参会队伍与经验分享者"],
        ["刘军教授 / 家畜生物学重点实验室", "西北农林科技大学家畜生物学重点实验室刘军教授"],
        ["西安动物医院", "西北农林科技大学西安动物医院的临床医生与院长"],
        ["猫咪驿站", "西安市浮生闲猫咪驿站工作人员"],
        ["张伟", "杨陵揉谷镇除张村羊场人员张伟"],
        ["杜欣愿", "武功县诚威奶山羊羊场负责人杜欣愿"],
        ["散养羊村民", "基层执业兽医与散养户"],
    ]
    add_table(doc, ["文件名/文本匹配", "完整 Stakeholder 标签"], rows, col_widths=[5.0, 8.5])

    add_heading(doc, "10.2 Stakeholder 类型推断规则（精确）", 2)
    rows_primary = [
        ["环境微生物 / ARG 专家", "钱勋、资环"],
        ["AI / 生物大数据专家", "赵天意"],
        ["湿实验 / 合成生物学专家", "罗自卫"],
        ["iGEM / Wiki 交流 stakeholder", "聂桓、交流会、iGEM"],
        ["动物健康 / 家畜专家", "刘军、家畜生物学"],
        ["兽医临床 stakeholder", "动物医院、医生、院长、检验中心"],
        ["养殖端 stakeholder", "羊场、养殖、散养户、村民、张伟、杜欣愿"],
        ["公众教育 stakeholder", "猫咪驿站、公众"],
    ]
    add_table(doc, ["类型标签", "匹配关键词（在 stakeholder 标签中）"], rows_primary,
              col_widths=[5.5, 7.5])

    add_heading(doc, "10.3 Stakeholder 类型推断规则（全文回退）", 2)
    rows_secondary = [
        ["环境微生物 / ARG 专家", "钱勋、环境、ARG、资环"],
        ["AI / 生物大数据专家", "赵天意、ESM、Oracle、模型、生物大数据"],
        ["湿实验 / 合成生物学专家", "罗自卫、湿实验、MIC、工程菌"],
        ["动物健康 / 家畜专家", "刘军、家畜、乳腺炎"],
        ["兽医临床 stakeholder", "动物医院、医生、院长、检验中心"],
        ["养殖端 stakeholder", "羊场、养殖、负责人、村民"],
        ["公众教育 stakeholder", "猫咪驿站、公众、工作人员"],
        ["iGEM / Wiki 交流 stakeholder", "交流会、Wiki、答辩、海报"],
    ]
    add_table(doc, ["类型标签", "匹配关键词（全文搜索）"], rows_secondary,
              col_widths=[5.5, 7.5])

    # ════════════════════════════════════════
    #  11. Graph parameters
    # ════════════════════════════════════════
    add_heading(doc, "11. 知识图谱参数（Φ₆ 图构建与 Φ₇ 图分析）", 1)

    add_heading(doc, "11.1 节点类型与信息层级", 2)
    rows = [[kind, str(level), LEVEL_LABELS[level]] for kind, level in NODE_LEVELS.items()]
    add_table(doc, ["节点类型 (kind)", "层级", "层级名"], rows, col_widths=[5.0, 2.0, 6.0])

    add_heading(doc, "11.2 图谱节点大小映射", 2)
    radius_ranges = [
        ["HP", "12–31 px"], ["Stakeholder", "13–34 px"], ["Feedback", "9–20 px"],
        ["Module", "9–25 px"], ["Action", "9–18 px"], ["Evidence", "6–11 px"],
        ["NextStep", "7–11 px"],
    ]
    add_table(doc, ["节点类型", "半径范围 (r_min + imp × (r_max − r_min))"], radius_ranges,
              col_widths=[5.0, 8.5])

    add_heading(doc, "11.3 图分析算法参数", 2)
    graph_algo_params = [
        ["PageRank 阻尼因子 α", "0.85"],
        ["PageRank 收敛阈值", "‖Δ‖∞ < 10⁻⁶"],
        ["社区发现算法", "greedy_modularity_communities (NetworkX)"],
    ]
    add_table(doc, ["参数", "取值"], graph_algo_params, col_widths=[5.5, 8.0])

    add_heading(doc, "11.4 图谱边关系类型", 2)
    edge_relations = [
        ["raised", "Stakeholder → Feedback", "利益相关者提出反馈"],
        ["recorded_in", "Feedback → HP", "反馈记录于某次 HP 循环"],
        ["led_to", "Feedback → Action", "反馈导致项目修改"],
        ["affects", "Feedback → Module", "反馈影响某项目模块"],
        ["supported_by", "Action → Evidence", "修改有证据支撑"],
        ["requires", "Action → NextStep", "需要后续步骤"],
        ["related_via_module", "Feedback ↔ Feedback", "共享模块的跨循环关联"],
    ]
    add_table(doc, ["关系类型", "方向", "含义"], edge_relations, col_widths=[4.0, 4.0, 5.5])

    add_heading(doc, "11.5 图谱主题指纹模式（12 类）", 2)
    add_para(doc, "用于 Feedback/Action 节点去重与主题归类。")
    for theme, keywords in THEME_PATTERNS.items():
        add_para(doc, f"【{theme}】{'、'.join(keywords)}")

    add_heading(doc, "11.6 摘要引擎评分词集", 2)
    add_para(doc, "强信号词（命中×3）：" + "、".join(STRONG_TERMS))
    add_para(doc, "实体词（命中×2）：" + "、".join(ENTITY_TERMS))

    add_heading(doc, "11.7 混合利益相关者排名权重（§7.3）", 2)
    add_para(doc, "拓扑得分 S_t(v) = 0.40·PR̃(v) + 0.35·C̃_D(v) + 0.25·C̃_B(v)")
    add_para(doc, "混合得分 S_h(v) = 0.30·S_t(v) + 0.35·P̄(v) + 0.25·m̄(v) + 0.10·σ̄ₑ(v)")
    hybrid_rows = [
        ["拓扑得分 S_t(v)", "0.30"],
        ["平均FCE优先级 P̄(v)", "0.35"],
        ["平均模块覆盖率 m̄(v)", "0.25"],
        ["平均证据强度 σ̄ₑ(v)", "0.10"],
    ]
    add_table(doc, ["分量", "权重"], hybrid_rows, col_widths=[5.5, 3.0])

    # ════════════════════════════════════════
    #  12. Sensitivity
    # ════════════════════════════════════════
    add_heading(doc, "12. 敏感性分析参数（§8）", 1)
    add_para(doc, "参数扰动范围：±20%。扰动四类参数：AHP权重（8个）、隶属函数边界参数（8个）、"
              "领域映射参数（14个）。度量：Spearman ρ、ΔP_max、成熟度跳变率。")
    rows = [[k, v] for k, v in FACTOR_LABELS.items()]
    add_table(doc, ["因素键", "中文标签"], rows, col_widths=[5.0, 7.0])

    # ════════════════════════════════════════
    #  13. Factor Computation Specs
    # ════════════════════════════════════════
    add_heading(doc, "13. 六因素计算公式规范（F₁-F₆）", 1)
    rows = [
        ["F₁ 闭环缺口", "F₁(ℓ) = (4−ℓ)/4", "ℓ ∈ {0,1,2,3,4}"],
        ["F₂ 跨模块影响", "F₂ = min(1, Σ_c μ_c / 4)", "μ_c 为 9 模块模糊隶属度"],
        ["F₃ 项目关键性", "F₃ = max_{c:μ_c>0} κ(c)·μ_c", "κ(c) 见 §2"],
        ["F₄ 时间紧迫度", "d≥120→0.20, 60≤d<120→0.45, 30≤d<60→0.70, d<30→1.00", "距 deadline 天数"],
        ["F₅ 证据不足度", "F₅ = 1 − σₑ", "σₑ 见 §5.1"],
        ["F₆ 利益相关者价值", "F₆ = max(0.50, max_{k⊑s} v(k))", "v(k) 见 §6"],
    ]
    add_table(doc, ["因素", "公式", "说明"], rows, col_widths=[3.0, 5.5, 5.0])

    # ════════════════════════════════════════
    #  14. Recommender Templates
    # ════════════════════════════════════════
    add_heading(doc, "14. 行动推荐器模板（Φ₆）", 1)

    add_heading(doc, "14.1 闭环层级→行动指令映射", 2)
    action_rows = [
        ["ℓ = 0-1", "—", "将反馈转化为具体项目修改，补充对应证据"],
        ["ℓ = 2", "—", "补充可展示证据（模型输出、实验数据、Wiki图表等）"],
        ["ℓ = 3", "r = 0", "将修改后材料返回利益相关者，完成二轮反馈"],
        ["ℓ ≥ 3", "r = 1", "整理为闭环案例研究，归档于 Wiki / 答辩"],
    ]
    add_table(doc, ["闭环层级", "回访标记", "行动指令"], action_rows, col_widths=[2.5, 2.5, 8.5])

    add_heading(doc, "14.2 模块化建议模板", 2)
    templates = [
        ("Model", "建议材料", "revised Field Score / model report, candidate ranking change table"),
        ("Model", "建议问题", "新的模型指标是否覆盖了 stakeholder 提出的关键判断？\n候选排序变化是否容易解释，是否存在误导风险？"),
        ("Software", "建议材料", "software panel screenshot, exported HP report card"),
        ("Software", "建议问题", "软件报告中的标签和证据等级是否足够清楚？"),
        ("Safety / Environment", "建议材料", "risk boundary wording draft, environmental / safety panel"),
        ("Safety / Environment", "建议问题", "风险措辞是否避免了过度承诺？\n哪些未来验证应被标注为优先？"),
        ("Material", "建议材料", "Evidence Matrix, wet-lab validation chain summary"),
        ("Material", "建议问题", "当前证据链是否足以支持候选筛选，而不是临床有效性声明？"),
        ("Implementation", "建议材料", "application scenario card, use-boundary checklist"),
        ("Implementation", "建议问题", "该应用场景下最容易被忽略的使用限制是什么？"),
        ("Education", "建议材料", "education material v2, public understanding feedback form"),
        ("Education", "建议问题", "受众是否能区分候选筛选、未来应用和现实治疗？"),
        ("Social Media", "建议材料", "HP timeline card, Stakeholder-Feedback-Action summary"),
        ("Social Media", "建议问题", "这段叙事是否清楚说明了谁改变了项目、改变了哪里？"),
    ]
    rows = [[m, t, c] for m, t, c in templates]
    add_table(doc, ["模块", "类型", "内容"], rows, col_widths=[3.0, 2.0, 8.5])

    # ════════════════════════════════════════
    #  15. Maturity FCE Signal Framework
    # ════════════════════════════════════════
    add_heading(doc, "15. 成熟度 FCE 信号变量框架与权重（Φ₅，模型规范 §6.1）", 1)
    add_para(doc, "每条 HP 反馈的六维成熟度通过 FCE 评估。对每个维度 i，计算若干信号变量 x_{i,j}，"
              "经单因素隶属函数 μ_{i,k}^{(j)}(x_{i,j}) 映射到各等级 k，"
              "综合隶属度 μ_{i,k} = Σ_j w_{i,j} · μ_{i,k}^{(j)}(x_{i,j})，Σ_j w_{i,j} = 1。"
              "参数由 3 位专家通过 Delphi 法标定（2轮，CV < 0.15），"
              "已完整实现在 config.py 和 maturity.py 中。")

    # Signal variable definitions
    signal_defs = [
        ["1. 设计反思迭代深度", "4",
         "x_{1,1} = ℓ/4\nx_{1,2} = |M∩D|/|D|, D={Model,Software,Problem Def}\n"
         "x_{1,3} = σ_e\nx_{1,4} = 设计反思关键词密度（饱和点=6）"],
        ["2. 真实场景探索", "4",
         "x_{2,1} = τ类型隶属度（医生0.95/企业0.80/NGO 0.65/学术0.40）\n"
         "x_{2,2} = |M∩R|/|R|, R={Implementation,Environment}\n"
         "x_{2,3} = 场景关键词密度（饱和点=6）\nx_{2,4} = σ_e"],
        ["3. 多元视角广度", "3",
         "x_{3,1} = min(1, Σ_c μ_c / 6)\n"
         "x_{3,2} = τ关键类型隶属度（UNIQUE→1.0, iGEM交流→0.65, 其他→0.50）\n"
         "x_{3,3} = F₂（跨模块影响因子，来自Φ₄）"],
        ["4. 正负影响预判", "5",
         "x_{4,1} = μ_Safety\nx_{4,2} = μ_Environment\n"
         "x_{4,3} = 风险关键词密度（饱和点=6）\n"
         "x_{4,4} = 缓解措施关键词密度（饱和点=4）\n"
         "x_{4,5} = 1（若命中风险证据产物）else 0"],
        ["5. 反馈落地响应", "4",
         "x_{5,1} = ℓ/4\nx_{5,2} = σ_e\n"
         "x_{5,3} = min(1, |a|/500)\nx_{5,4} = r（0或1）"],
        ["6. 局限性坦诚认知", "4",
         "x_{6,1} = min(1, λ/8), λ=局限性词命中数\n"
         "x_{6,2} = μ_Safety\nx_{6,3} = 边界语言密度（饱和点=3）\n"
         "x_{6,4} = 1（若命中边界文档产物）else 0"],
    ]
    add_table(doc, ["维度 i", "信号数", "信号变量 x_{i,j}"], signal_defs,
              col_widths=[3.5, 1.2, 9.0])

    # ════════════════════════════════════════
    #  16. Maturity Signal Weights
    # ════════════════════════════════════════
    add_heading(doc, "16. 成熟度信号权重 w_{i,j}（§6.1）", 1)
    add_para(doc, "每个维度内的信号权重，满足 Σ_j w_{i,j} = 1。"
              "由 3 位专家通过 Delphi 法标定（2轮，CV < 0.15）。")

    for dim_key in DIMENSION_ORDER:
        idx = DIMENSION_ORDER.index(dim_key) + 1
        short = DIMENSION_SHORT[dim_key]
        weights = MATURITY_SIGNAL_WEIGHTS[dim_key]
        n = len(weights)

        add_heading(doc, f"16.{idx} {short}（{n} 个信号）", 2)
        headers = [f"x_{{{idx},{j+1}}}" for j in range(n)]
        row = [f"{w:.2f}" for w in weights]
        add_table(doc, ["信号变量"] + headers,
                  [[f"权重 w_{{{idx},j}}"] + row],
                  col_widths=[3.5] + [2.5] * n)

    # ════════════════════════════════════════
    #  17. Maturity Membership Functions
    # ════════════════════════════════════════
    add_heading(doc, "17. 成熟度隶属函数参数 μ_{i,k}^{(j)} (a, b, c, d)（§6.1）", 1)
    add_para(doc, "六组标准梯形隶属函数（每等级一个），定义在信号值域 [0,1] 上。"
              "参数由 Delphi 法标定（3位专家，2轮，CV < 0.15）。\n"
              "各等级函数类型自动检测：a=b=0→降半梯形；c=d=1→升半梯形；否则→三角/梯形。\n"
              "满足：(1) 覆盖性 Σ_k μ_k(x)≈1；(2) 凸性；(3) 平滑过渡——相邻等级在交叉区重叠。")

    shape_types = [
        "降半梯形",   # m0
        "三角形",     # m1
        "三角形",     # m2
        "三角形",     # m3
        "三角形",     # m4
        "升半梯形",   # m5
    ]
    shape_descs = [
        "x≤0→0.5(软底), 线性降至 x≥0.26→0",
        "0→升至0.26→1, 0.34→开始降, 0.48→0",
        "0→升至0.48→1, 0.56→开始降, 0.70→0",
        "0→升至0.70→1, 0.78→开始降, 0.90→0",
        "0→升至0.90→1, 0.94→开始降, 0.98→0",
        "x≤0.92→0, 线性升至 x≥0.97→1",
    ]

    m_rows = []
    for k in range(6):
        a, b, c, d = MATURITY_MEMBERSHIP_PARAMS[f"m{k}"]
        m_rows.append([
            f"m{k}（等级 {k}）", shape_types[k],
            f"{a:.2f}", f"{b:.2f}", f"{c:.2f}", f"{d:.2f}",
            shape_descs[k],
        ])
    add_table(doc, ["参数键", "形状", "a", "b", "c", "d", "形状描述"],
              m_rows, col_widths=[2.5, 1.8, 1.2, 1.2, 1.2, 1.2, 5.0])

    # ════════════════════════════════════════
    #  18. Maturity FCE Synthesis & Level Determination
    # ════════════════════════════════════════
    add_heading(doc, "18. 成熟度 FCE 合成方法与等级判定（§6.2）", 1)

    add_heading(doc, "18.1 合成算子", 2)
    add_para(doc, "采用 M(·,+) 加权平均算子（与优先级 FCE 一致）："
              "μ_{i,k} = Σ_j w_{i,j} · μ_{i,k}^{(j)}(x_{i,j})。"
              "合成后归一化至 Σ_k μ_{i,k} = 1。")

    add_heading(doc, "18.2 等级判定规则", 2)
    add_para(doc, "规则1（最大隶属度原则 — 占优情形）：若 ∃k: μ_{i,k} > 0.5，"
              "则 m_i = argmax_k μ_{i,k}。此时 max 是唯一的（定理6.1："
              "若 μ_k > 0.5 则 Σ_{j≠k} μ_j < 0.5，不存在 j≠k 满足 μ_j ≥ μ_k）。\n"
              "规则2（级别特征值法 — 所有情形均计算）："
              "m_i* = Σ(k·μ_k^γ)/Σ(μ_k^γ)，γ=2。"
              "m_i* ∈ [0,5] 为连续成熟度得分，保留梯度信息供跨卡片比较。")

    add_heading(doc, "18.3 级别特征值参数 γ", 2)
    gamma_rows = [
        ["γ", "2", "幂加权指数 — 放大高隶属度等级贡献，抑制低隶属度噪声"],
        ["γ = 1", "退化为普通加权平均", "对噪声敏感，所有等级等权重贡献"],
        ["γ → ∞", "退化为最大隶属度原则", "丢失梯度信息"],
        ["m_i* 值域", "[0, 5]", "2.7 = 稳固 L2 接近 L3；2.1 = 刚达 L2"],
    ]
    add_table(doc, ["参数", "取值", "说明"], gamma_rows, col_widths=[3.5, 4.0, 6.0])

    # ════════════════════════════════════════
    #  19. Evidence Extractor Supplementary
    # ════════════════════════════════════════
    add_heading(doc, "19. 证据提取器补充参数（Φ₁）", 1)

    add_heading(doc, "19.1 命名项目产物列表", 2)
    add_para(doc, "独立扫描以下命名产物，命中即追加到证据列表：")
    named_artifacts = [
        "PDES", "Environmental Degradation Panel", "Field Score",
        "Evidence Matrix", "TAM-Flow", "Risk Boundary Panel", "Oracle", "RAFT",
    ]
    add_para(doc, "、".join(named_artifacts))

    add_heading(doc, "19.2 强证据信号词", 2)
    add_para(doc, "命中以下词时，访谈记录仅作为辅助证据：")
    add_para(doc, "MIC、溶血、CCK、TEM、MD、质谱、PDES、TAM-Flow")

    add_heading(doc, "19.3 证据去重规则", 2)
    dedup_rules = [
        ["CCK → CCK-8", "CCK-8 更精确"],
        ["MD → 分子动力学", "分子动力学更完整"],
    ]
    add_table(doc, ["规则", "理由"], dedup_rules, col_widths=[5.0, 8.5])

    add_heading(doc, "19.4 段落标记", 2)
    markers = [
        ["初始问题（≤520字符）", "为什么：、为什么:、为什么、设计问题：、设计需要被验证的问题："],
        ["核心反馈（≤900字符）", "我们学到了什么：、我们学到了什么:、我们学到了什么、专家反馈、学到了什么：、反馈"],
        ["项目修改（≤900字符）", "我们如何修改项目：、我们如何修改项目、AMPlify 的修改、形成 v2、Learn、如何改变"],
        ["段落终止", "了解更多、本轮 HP、这次调研为什么、为什么这一站、图 1、仍需谨慎"],
    ]
    add_table(doc, ["字段", "起始/终止标记"], markers, col_widths=[4.0, 9.5])

    # ════════════════════════════════════════
    #  20. Global Constants
    # ════════════════════════════════════════
    add_heading(doc, "20. 全局常量与归一化参数", 1)

    add_heading(doc, "20.1 证据强度", 2)
    const_rows = [
        ["证据强度基线值", "0.35", "未匹配到证据关键词的默认分"],
        ["证据强度计算", "σ_e = (1/|E|) Σ_i σ(e_i)", "E=∅ 时 σ_e=0"],
    ]
    add_table(doc, ["参数", "取值", "说明"], const_rows, col_widths=[4.5, 5.0, 4.0])

    add_heading(doc, "20.2 成熟度信号归一化参数", 2)
    norm_rows = [
        ["设计反思关键词饱和点", str(_DESIGN_KW_SATURATION), "x_{1,4} = min(1, hits/6)"],
        ["场景关键词饱和点", str(_SCENE_KW_SATURATION), "x_{2,3} = min(1, hits/6)"],
        ["风险关键词饱和点", str(_RISK_KW_SATURATION), "x_{4,3} = min(1, hits/6)"],
        ["缓解关键词饱和点", str(_MITIGATION_KW_SATURATION), "x_{4,4} = min(1, hits/4)"],
        ["边界语言饱和点", str(_BOUNDARY_KW_SATURATION), "x_{6,3} = min(1, hits/3)"],
        ["局限性词计数分母", str(_LIMITATION_KW_DENOM), "x_{6,1} = min(1, λ/8)"],
        ["行动文本长度分母", str(_ACTION_LEN_DENOM), "x_{5,3} = min(1, |a|/500)"],
        ["加权模块数分母", str(_MODULE_COUNT_DENOM), "x_{3,1} = min(1, Σμ_c/6)"],
        ["零信号处理", "均匀先验 1/6", "x≈0 时各等级均分贡献，避免偏向L0"],
    ]
    add_table(doc, ["参数", "取值", "对应信号"], norm_rows, col_widths=[4.5, 3.0, 6.0])

    add_heading(doc, "20.3 FCE 合成参数", 2)
    fce_rows = [
        ["模糊合成算子", "M(·,+) 加权平均型"],
        ["去模糊化方法（优先级）", "重心法 P = Σ bⱼ·cⱼ"],
        ["等级判定（成熟度）", "最大隶属度 + 级别特征值 γ=2"],
        ["零信号处理（成熟度 v3.1）", "x≈0 → 均匀 1/6 先验，避免稀疏信号拖向 L0"],
        ["归一化约束", "Σ_j b_j = 1，软件归一化吸收浮点误差"],
    ]
    add_table(doc, ["参数", "取值"], fce_rows, col_widths=[5.0, 8.5])

    # ════════════════════════════════════════
    #  21. Parameter Completeness Check
    # ════════════════════════════════════════
    add_heading(doc, "21. 参数完整性清单", 1)
    add_para(doc, "以下列出模型规范中定义的全部可配置参数及其在代码中的实现位置。")

    param_checklist = [
        ["模块关键词集 K_c（9个模块）", "config.py → CATEGORIES"],
        ["模块隶属度阈值 (α_c, β_c)（9对）", "config.py → MODULE_THRESHOLDS"],
        ["模块关键性 κ(c)（9个）", "config.py → MODULE_CRITICALITY"],
        ["AHP 判断矩阵 J₁, J₂, J（3个）", "config.py → J1/J2/J2ND_MATRIX"],
        ["AHP 权重 A₁, A₂, A（8个）", "config.py → A1/A2/A2ND_WEIGHTS"],
        ["FCE 量化向量 C（4个）", "config.py → QUANT_VECTOR"],
        ["FCE 隶属函数参数 (a,b,c,d)（4组）", "config.py → MEMBERSHIP_PARAMS"],
        ["成熟度信号权重 w_{i,j}（24个）", "config.py → MATURITY_SIGNAL_WEIGHTS"],
        ["成熟度隶属函数参数 (a,b,c,d)（6组）", "config.py → MATURITY_MEMBERSHIP_PARAMS"],
        ["成熟度信号归一化常数（8个）", "maturity.py → _*_SATURATION/_DENOM 常量"],
        ["证据关键词权重（17个）", "config.py → EVIDENCE_KEYWORDS"],
        ["证据屏蔽词（5个）", "config.py → EVIDENCE_BLOCKED"],
        ["行动关键词（17个）", "config.py → ACTION_KEYWORDS"],
        ["返回关键词（6个）", "config.py → RETURN_KEYWORDS"],
        ["Stakeholder 关键词价值（9个）", "config.py → STAKEHOLDER_VALUE_KEYWORDS"],
        ["Stakeholder 类型价值（9个）", "config.py → STAKEHOLDER_TYPE_VALUES"],
        ["成熟度锚定描述（6维×6级）", "maturity.py → LEVEL_ANCHORS"],
        ["成熟度信号关键词（14类）", "maturity.py → 各关键词列表"],
        ["闭环状态机 L0-L4", "status.py → STATUS_NAMES"],
        ["Stakeholder 映射表（11条）", "extractor.py → known_stakeholders"],
        ["Stakeholder 类型推断规则（16条）", "extractor.py → stakeholder_rules"],
        ["Stakeholder 类型→成熟度信号映射", "maturity.py → _stakeholder_type_membership"],
        ["图谱主题指纹模式（12类）", "graph_builder.py → THEME_PATTERNS"],
        ["图谱节点层级定义（7种）", "graph_builder.py → NODE_LEVELS"],
        ["PageRank α=0.85", "graph_builder.py → compute_pagerank"],
        ["混合排名权重（4个分量）", "graph_builder.py → top_stakeholders_hybrid"],
    ]
    add_table(doc, ["参数项", "源文件位置"], param_checklist, col_widths=[7.0, 6.5])

    # ════════════════════════════════════════
    #  Save
    # ════════════════════════════════════════
    output_path = ROOT / "HP_Compass_模型白盒参数总表_v3.docx"
    doc.save(str(output_path))
    print(f"Saved: {output_path}")
    print(f"File size: {output_path.stat().st_size / 1024:.1f} KB")


# ════════════════════════════════════════════
#  Utilities
# ════════════════════════════════════════════

def _chunk_list(items: list[str], n_cols: int) -> list[list[str]]:
    result = []
    for i in range(0, len(items), n_cols):
        row = []
        for j in range(n_cols):
            idx = i + j
            if idx < len(items):
                row.append(str(idx + 1))
                row.append(items[idx])
            else:
                row.append("")
                row.append("")
        result.append(row)
    return result


def _add_matrix_table(doc, matrix, labels):
    n = len(labels)
    headers = [""] + labels
    rows = []
    for i, label in enumerate(labels):
        row = [label]
        for j in range(n):
            val = matrix[i][j]
            if val >= 1:
                row.append(f"{val:.0f}")
            else:
                row.append(f"1/{int(1/val)}")
        rows.append(row)
    add_table(doc, headers, rows, col_widths=[3.5] + [3.0] * n)


if __name__ == "__main__":
    main()
